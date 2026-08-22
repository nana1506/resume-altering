import io
import re
from typing import List, Literal, Optional, Tuple, Dict, Any
from pydantic import BaseModel, Field

class CVEntry(BaseModel):
    type: Literal["bullet", "paragraph", "subheading", "skill_line"]
    text: str
    label: Optional[str] = None

class CVSection(BaseModel):
    name: str  # e.g. "Profile Summary", "Work Experience", "Skills & Technologies", "Education"
    entries: List[CVEntry] = Field(default_factory=list)

class ParsedCV(BaseModel):
    name: str = ""
    headline: str = ""
    contact: str = ""
    sections: List[CVSection] = Field(default_factory=list)

# Known section header aliases mapped to canonical display names
CANONICAL_SECTION_MAP = {
    "summary": "Profile Summary",
    "profile summary": "Profile Summary",
    "professional summary": "Profile Summary",
    "executive summary": "Profile Summary",
    "about": "Profile Summary",
    "about me": "Profile Summary",
    "work experience": "Work Experience",
    "experience": "Work Experience",
    "professional experience": "Work Experience",
    "employment history": "Work Experience",
    "work history": "Work Experience",
    "career history": "Work Experience",
    "skills": "Skills & Technologies",
    "technical skills": "Skills & Technologies",
    "core competencies": "Skills & Technologies",
    "skills & tools": "Skills & Technologies",
    "technologies": "Skills & Technologies",
    "key skills": "Skills & Technologies",
    "education": "Education",
    "academic background": "Education",
    "qualifications": "Education",
    "certifications": "Certifications",
    "certificates": "Certifications",
    "licenses & certifications": "Certifications",
    "projects": "Projects",
    "key projects": "Projects",
    "personal projects": "Projects",
    "selected projects": "Projects",
    "languages": "Languages",
    "awards": "Awards & Honors",
    "honors": "Awards & Honors",
    "publications": "Publications",
    "volunteer": "Volunteer Experience",
    "volunteering": "Volunteer Experience",
    "additional": "Additional Information",
    "additional information": "Additional Information",
}

def clean_text(text: str) -> str:
    """Normalizes whitespace and removes unwanted artifacts."""
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.replace('\xa0', ' ').replace('\u200b', '')
    return text.strip()

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extracts plain text from a PDF file using pdfplumber without table duplication.
    extract_text() naturally captures table contents as flowing text; we do not blindly append extract_tables().
    """
    import pdfplumber

    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_chunks.append(page_text.strip())
            else:
                # Fallback to tables only if extract_text failed or was empty on this page
                tables = page.extract_tables() or []
                table_lines = []
                for table in tables:
                    for row in table:
                        filtered_row = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                        if filtered_row:
                            table_lines.append(" | ".join(filtered_row))
                if table_lines:
                    text_chunks.append("\n".join(table_lines))

    extracted = "\n\n".join(text_chunks).strip()
    return clean_text(extracted)

def parse_docx(file_bytes: bytes) -> str:
    """
    Extracts plain text from a DOCX file using python-docx.
    Appends paragraph text and only non-redundant table text.
    """
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    para_chunks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            para_chunks.append(text)

    # In DOCX, table cells are separate from doc.paragraphs
    table_chunks = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            unique_cells = []
            for c in row_cells:
                if not unique_cells or unique_cells[-1] != c:
                    unique_cells.append(c)
            if unique_cells:
                table_chunks.append(" | ".join(unique_cells))

    all_chunks = para_chunks
    if table_chunks:
        para_blob = "\n".join(para_chunks)
        new_table_chunks = [t for t in table_chunks if t not in para_blob]
        if new_table_chunks:
            all_chunks.append("\n".join(new_table_chunks))

    extracted = "\n\n".join(all_chunks).strip()
    return clean_text(extracted)

def is_section_header(line: str) -> Optional[str]:
    """
    Checks if a line represents a CV section header.
    Returns the normalized canonical section name if true, or None.
    """
    clean = line.strip()
    if not clean or len(clean) > 45:
        return None

    clean_lower = clean.lower().rstrip(":-").strip()

    # Direct canonical lookup
    if clean_lower in CANONICAL_SECTION_MAP:
        return CANONICAL_SECTION_MAP[clean_lower]

    # Starts with known header prefix
    for prefix, canonical in CANONICAL_SECTION_MAP.items():
        if clean_lower == prefix or clean_lower.startswith(f"{prefix}:"):
            return canonical

    # Uppercase header candidate (e.g. "EXPERIENCE", "PROJECTS", "SKILLS")
    if clean.isupper() and 3 <= len(clean) <= 30 and not any(c in clean for c in ["@", "http", "|", "/", "\\"]):
        return clean.title()

    # Ends with colon candidate (e.g. "Work Experience:")
    if clean.endswith(":") and len(clean) <= 35 and not clean.startswith("http"):
        base = clean.rstrip(":").strip()
        if base.lower() in CANONICAL_SECTION_MAP:
            return CANONICAL_SECTION_MAP[base.lower()]
        return base.title()

    return None

def is_contact_line(line: str) -> bool:
    """Detects if a line contains contact details (email, phone, location, links)."""
    lower = line.lower()
    has_email = "@" in lower and "." in lower
    has_link = any(kw in lower for kw in ["linkedin.com", "github.com", "http://", "https://", "www."])
    has_phone = bool(re.search(r'(\+\d{1,3}[\s-]?)?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{3,4}', line))
    has_pipe_delimiters = line.count("|") >= 2
    return (has_email or has_link or has_phone or has_pipe_delimiters) and len(line) < 160

def is_date_or_subheading(line: str) -> bool:
    """Detects job titles, companies, or date ranges within experience/education."""
    # Date range pattern: e.g. "2020 - 2023", "Jan 2021 - Present", "2019 – Present", "05/2020 - 08/2022"
    date_pattern = r'(?:\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)\b\.?\s*)?\b\d{4}\b\s*(?:-|–|—|to)\s*(?:\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)\b\.?\s*)?(?:\b\d{4}\b|\bPresent\b|\bCurrent\b)'
    if re.search(date_pattern, line, re.IGNORECASE):
        return True

    # Role at Company patterns (e.g. "Senior Engineer at Google", "Product Designer | Figma")
    if (" at " in line or " | " in line or " - " in line) and len(line) < 90:
        if not line.startswith("•") and not line.startswith("-") and not line.startswith("*"):
            return True

    return False

def parse_skill_line(line: str) -> Optional[Tuple[str, str]]:
    """
    Checks if a line in a Skills section matches a category label + items pattern.
    Returns (label, values) or None.
    """
    clean = line.strip()
    if not clean:
        return None

    # Strip bullet marker if present
    bullet_prefix = re.match(r'^[•\-\*▪▫–—\u2022\u25cf\u25cb\u25aa\u25ab]\s*', clean)
    if bullet_prefix:
        clean = clean[bullet_prefix.end():].strip()

    # 1. Pattern: "Category Label: Item1, Item2, Item3" or "Category Label | Item1, Item2" or "Category Label – Item1..."
    sep_match = re.match(r'^([A-Za-z0-9\s&/\(\)\+]{2,35}?)\s*[:|–—]\s*(.+)$', clean)
    if sep_match:
        label = sep_match.group(1).strip()
        values = sep_match.group(2).strip()
        if len(label.split()) <= 5 and not label.endswith('.'):
            return label, values

    # 2. Pattern: "Technical Skills · Python · SQL · React"
    dot_match = re.match(r'^([A-Za-z0-9\s&/\(\)\+]{2,30}?)\s*[·•▪]\s*(.+)$', clean)
    if dot_match:
        label = dot_match.group(1).strip()
        values = dot_match.group(2).strip()
        if len(label.split()) <= 4:
            return label, values

    # 3. Known category keywords at start of line
    known_skill_prefixes = [
        "technical skills", "hard skills", "soft skills", "core competencies",
        "key skills", "programming languages", "languages", "frameworks",
        "tools & technologies", "tools", "databases", "cloud & devops",
        "bi & visualization", "visualization", "methodologies", "analytics",
        "certifications", "domain knowledge", "professional skills"
    ]
    lower = clean.lower()
    for prefix in known_skill_prefixes:
        if lower.startswith(prefix) and len(clean) > len(prefix) + 2:
            remainder = clean[len(prefix):].lstrip(" :|–—·•-").strip()
            if remainder:
                return clean[:len(prefix)].strip(), remainder

    return None

def structure_cv_text(plain_text: str) -> ParsedCV:
    """
    Transforms extracted CV text into a clean, structured ParsedCV model.
    Merges line-wrapped continuation lines and categorizes skills entries cleanly.
    """
    lines = [l.strip() for l in plain_text.strip().split("\n") if l.strip()]
    if not lines:
        return ParsedCV(name="Candidate", headline="", contact="", sections=[])

    # 1. Header extraction (Name, Headline, Contact)
    idx = 0
    name = lines[0]
    headline = ""
    contact = ""

    idx = 1
    while idx < min(6, len(lines)):
        candidate_line = lines[idx]
        if is_section_header(candidate_line):
            break

        if is_contact_line(candidate_line):
            if not contact:
                contact = candidate_line
            else:
                contact += f" | {candidate_line}"
            idx += 1
            continue

        # If not contact and not section header, it is likely a professional headline
        if not headline and len(candidate_line) < 100:
            headline = candidate_line
            idx += 1
            continue

        break

    # 2. Section and Entry parsing
    sections: List[CVSection] = []
    current_section: Optional[CVSection] = None

    bullet_regex = re.compile(r'^[•\-\*▪▫–—\u2022\u25cf\u25cb\u25aa\u25ab]\s*|\s*^\d+\.\s*')

    while idx < len(lines):
        line = lines[idx]
        section_name = is_section_header(line)

        if section_name:
            # Check if this section already exists to avoid duplicates
            existing = next((s for s in sections if s.name.lower() == section_name.lower()), None)
            if existing:
                current_section = existing
            else:
                current_section = CVSection(name=section_name, entries=[])
                sections.append(current_section)
            idx += 1
            continue

        # If we haven't encountered a section header yet, default to Profile Summary
        if current_section is None:
            current_section = CVSection(name="Profile Summary", entries=[])
            sections.append(current_section)

        is_skills_sec = any(
            kw in current_section.name.lower()
            for kw in ["skill", "technolog", "competenc", "tool"]
        )

        # Fix 1 & Fix 3: Check for Skill Category Line
        if is_skills_sec:
            skill_parsed = parse_skill_line(line)
            if skill_parsed:
                lbl, val = skill_parsed
                current_section.entries.append(CVEntry(type="skill_line", label=lbl, text=val))
                idx += 1
                continue

        # Fix 1: Check if this line is a continuation of the previous entry
        if current_section.entries:
            prev_entry = current_section.entries[-1]
            prev_text = prev_entry.text.strip()
            
            # Sentence terminal check
            is_terminal = prev_text.endswith(('.', '!', '?', ':'))
            is_new_bullet = bool(bullet_regex.match(line))
            is_new_section = bool(is_section_header(line))
            is_subhead = is_date_or_subheading(line)
            is_skill_line = is_skills_sec and bool(parse_skill_line(line))

            # If previous line was cut off mid-sentence and this is not a distinct new item
            if (
                not is_terminal
                and not is_new_bullet
                and not is_new_section
                and not is_subhead
                and not is_skill_line
                and prev_entry.type in ("bullet", "paragraph", "skill_line")
            ):
                # Clean merge into previous entry
                merged_text = f"{prev_text} {line}".strip()
                prev_entry.text = re.sub(r'\s+', ' ', merged_text)
                idx += 1
                continue

        # Check standard entry types
        if bullet_regex.match(line):
            clean_bullet = bullet_regex.sub('', line).strip()
            if clean_bullet:
                current_section.entries.append(CVEntry(type="bullet", text=clean_bullet))
        elif is_date_or_subheading(line) and "experience" in current_section.name.lower():
            current_section.entries.append(CVEntry(type="subheading", text=line))
        elif len(line) < 80 and not line.endswith(".") and ("experience" in current_section.name.lower() or "education" in current_section.name.lower()) and idx + 1 < len(lines) and bullet_regex.match(lines[idx + 1]):
            current_section.entries.append(CVEntry(type="subheading", text=line))
        else:
            # Paragraph entry (e.g. summary or inline text)
            current_section.entries.append(CVEntry(type="paragraph", text=line))

        idx += 1

    return ParsedCV(
        name=name,
        headline=headline,
        contact=contact,
        sections=sections
    )

def cv_to_plain_text(cv: ParsedCV) -> str:
    """Converts a ParsedCV object back to formatted text for LLM prompt context."""
    chunks = []
    if cv.name:
        chunks.append(cv.name)
    if cv.headline:
        chunks.append(cv.headline)
    if cv.contact:
        chunks.append(cv.contact)
    if chunks:
        chunks.append("-" * 30)

    for sec in cv.sections:
        chunks.append(f"\n[{sec.name.upper()}]")
        for idx, entry in enumerate(sec.entries):
            if entry.type == "skill_line":
                label_str = f"{entry.label}: " if entry.label else ""
                chunks.append(f"  • [entry {idx}]: {label_str}{entry.text}")
            elif entry.type == "bullet":
                chunks.append(f"  • [entry {idx}]: {entry.text}")
            elif entry.type == "subheading":
                chunks.append(f"  ### [entry {idx}]: {entry.text}")
            else:
                chunks.append(f"  [entry {idx}]: {entry.text}")

    return "\n".join(chunks)

def parse_file_to_structured_cv(filename: str, file_bytes: bytes) -> Tuple[str, ParsedCV, str]:
    """
    Parses a document file and returns (plain_text, parsed_cv_structure, content_type).
    """
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        raw_text = parse_pdf(file_bytes)
        content_type = "application/pdf"
    elif lower_filename.endswith(".docx"):
        raw_text = parse_docx(file_bytes)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif lower_filename.endswith(".txt"):
        raw_text = clean_text(file_bytes.decode("utf-8", errors="ignore"))
        content_type = "text/plain"
    else:
        raise ValueError(f"Unsupported file format: {filename}. Please upload a PDF or DOCX file.")

    structured_cv = structure_cv_text(raw_text)
    return raw_text, structured_cv, content_type

def extract_text_from_file(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    """
    Backward-compatible wrapper returning (plain_text, content_type).
    """
    raw_text, _, content_type = parse_file_to_structured_cv(filename, file_bytes)
    return raw_text, content_type
